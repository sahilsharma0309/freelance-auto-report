"""Branded Word (.docx) export via python-docx.

Mirrors the PDF report: branded header, chart images + insights, tables,
and a "Prepared by ..." watermark with page numbers in the footer.
"""

import io
from datetime import date

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from core.analysis import AnalysisResult
from docx.oxml import parse_xml

from core.branding import (
    ACCENT_COLOR,
    BRAND_NAME,
    LOGO_PATH,
    PAGE_WATERMARK_TEXT,
    PRIMARY_COLOR,
    SIGNATURE_PATH,
    WATERMARK_TEXT,
)

# Classic Word text watermark: a rotated VML shape anchored in the header,
# centered on the page behind the body text.
_WATERMARK_PICT = (
    '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:pPr/><w:r><w:pict xmlns:v="urn:schemas-microsoft-com:vml" '
    'xmlns:o="urn:schemas-microsoft-com:office:office">'
    '<v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" '
    'adj="10800" path="m@7,l@8,m@5,21600l@6,21600e">'
    '<v:formulas><v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/>'
    '<v:f eqn="sum 21600 0 @1"/><v:f eqn="sum 0 0 @2"/><v:f eqn="sum 21600 0 @3"/>'
    '<v:f eqn="if @0 @3 0"/><v:f eqn="if @0 21600 @1"/><v:f eqn="if @0 0 @2"/>'
    '<v:f eqn="if @0 @4 21600"/><v:f eqn="mid @5 @6"/><v:f eqn="mid @8 @5"/>'
    '<v:f eqn="mid @7 @8"/><v:f eqn="mid @6 @7"/><v:f eqn="sum @6 0 @5"/></v:formulas>'
    '<v:path textpathok="t" o:connecttype="custom" o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800"/>'
    '<v:textpath on="t" fitshape="t"/><o:lock v:ext="edit" text="t" shapetype="t"/>'
    '</v:shapetype>'
    '<v:shape id="BrandWatermark" type="#_x0000_t136" '
    'style="position:absolute;margin-left:0;margin-top:0;width:480pt;height:120pt;'
    'rotation:315;z-index:-251654144;mso-position-horizontal:center;'
    'mso-position-horizontal-relative:margin;mso-position-vertical:center;'
    'mso-position-vertical-relative:margin" '
    f'o:allowincell="f" fillcolor="{PRIMARY_COLOR}" stroked="f">'
    '<v:fill opacity="0.05"/>'
    '<v:textpath style="font-family:&quot;Helvetica&quot;;font-size:1pt" '
    f'string="{PAGE_WATERMARK_TEXT}"/>'
    '</v:shape></w:pict></w:r></w:p>'
)


def _add_page_watermark(document: Document) -> None:
    header = document.sections[0].header
    header._element.append(parse_xml(_WATERMARK_PICT))

MAX_TABLE_ROWS = 20

PRIMARY_RGB = RGBColor.from_string(PRIMARY_COLOR.lstrip("#"))
ACCENT_RGB = RGBColor.from_string(ACCENT_COLOR.lstrip("#"))


def _add_bottom_border(paragraph, hex_color: str) -> None:
    """Thin colored rule under a paragraph (used for the branded header)."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), hex_color.lstrip("#"))
    borders.append(bottom)
    p_pr.append(borders)


def _add_page_number(paragraph) -> None:
    """Append a live PAGE field to a paragraph."""
    run = paragraph.add_run()
    for element, text in (("w:fldChar", None), ("w:instrText", "PAGE"), ("w:fldChar", None)):
        node = OxmlElement(element)
        if element == "w:instrText":
            node.set(qn("xml:space"), "preserve")
            node.text = f" {text} "
        run._r.append(node)
    run._r[0].set(qn("w:fldCharType"), "begin")
    run._r[-1].set(qn("w:fldCharType"), "end")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x8A, 0x8F, 0x98)


def _shade_cell(cell, hex_color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), hex_color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shading)


def _build_header(document: Document) -> None:
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    if LOGO_PATH.exists():
        paragraph.add_run().add_picture(str(LOGO_PATH), height=Inches(0.45))
        paragraph.add_run("   ")
    brand_run = paragraph.add_run(BRAND_NAME)
    brand_run.font.bold = True
    brand_run.font.size = Pt(13)
    brand_run.font.color.rgb = PRIMARY_RGB
    _add_bottom_border(paragraph, ACCENT_COLOR)


def _build_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mark = paragraph.add_run(f"{WATERMARK_TEXT}   ·   Page ")
    mark.font.italic = True
    mark.font.size = Pt(8)
    mark.font.color.rgb = ACCENT_RGB
    _add_page_number(paragraph)


def _add_table(document: Document, df: pd.DataFrame) -> None:
    shown = df.head(MAX_TABLE_ROWS)
    table = document.add_table(rows=1, cols=len(shown.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(shown.columns):
        cell = table.rows[0].cells[idx]
        cell.text = str(column)
        _shade_cell(cell, PRIMARY_COLOR)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
    for _, row in shown.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = "" if pd.isna(value) else str(value)
            for run in cells[idx].paragraphs[0].runs:
                run.font.size = Pt(9)
    if len(df) > MAX_TABLE_ROWS:
        note = document.add_paragraph(f"Showing first {MAX_TABLE_ROWS} of {len(df)} rows.")
        note.runs[0].font.italic = True
        note.runs[0].font.size = Pt(8)


def _add_section(document: Document, result: AnalysisResult) -> None:
    question = document.add_paragraph()
    q_run = question.add_run(result.question)
    q_run.font.bold = True
    q_run.font.size = Pt(12)
    q_run.font.color.rgb = PRIMARY_RGB

    if result.kind == "chart":
        if result.chart_path:
            document.add_picture(result.chart_path, width=Inches(6.0))
        if result.text:
            insight = document.add_paragraph(result.text)
            insight.runs[0].font.size = Pt(10)
    elif result.kind == "dataframe" and result.dataframe is not None:
        _add_table(document, result.dataframe)
    elif result.kind == "error":
        error = document.add_paragraph(result.text)
        error.runs[0].font.color.rgb = RGBColor(0xAA, 0x33, 0x33)
    else:
        body = document.add_paragraph(result.text)
        body.runs[0].font.size = Pt(10)
    document.add_paragraph()


def _add_summary(document: Document, results: list[AnalysisResult]) -> None:
    insights = [r.text for r in results if r.kind == "chart" and r.text]
    if len(insights) < 2:
        return
    heading = document.add_paragraph()
    run = heading.add_run("Key Findings")
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = PRIMARY_RGB
    for text in insights[:6]:
        bullet = document.add_paragraph(text, style="List Bullet")
        for run in bullet.runs:
            run.font.size = Pt(9.5)
    document.add_paragraph()


def _add_signature(document: Document) -> None:
    block = document.add_paragraph()
    block.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if SIGNATURE_PATH.exists():
        block.add_run().add_picture(str(SIGNATURE_PATH), height=Inches(0.5))
        block.add_run("\n")
    name = block.add_run(BRAND_NAME)
    name.font.size = Pt(9.5)
    name.font.bold = True
    name.font.color.rgb = PRIMARY_RGB
    date_run = block.add_run(f"\n{date.today():%d %B %Y}")
    date_run.font.size = Pt(8)
    date_run.font.color.rgb = RGBColor(0x8A, 0x8F, 0x98)


def _add_kpi_row(document: Document, kpis) -> None:
    table = document.add_table(rows=1, cols=len(kpis))
    table.style = "Table Grid"
    for idx, kpi in enumerate(kpis):
        cell = table.rows[0].cells[idx]
        paragraph = cell.paragraphs[0]
        label = paragraph.add_run(kpi.label.upper() + "\n")
        label.font.size = Pt(7.5)
        label.font.color.rgb = RGBColor(0x6A, 0x70, 0x7A)
        value = paragraph.add_run(kpi.value)
        value.font.bold = True
        value.font.size = Pt(14)
        value.font.color.rgb = PRIMARY_RGB
        if kpi.delta:
            delta = paragraph.add_run("\n" + kpi.delta)
            delta.font.size = Pt(7.5)
            delta.font.color.rgb = ACCENT_RGB
    document.add_paragraph()


def export_docx(results: list[AnalysisResult], title: str, dataset_name: str,
                kpis=None, client_name: str = "") -> bytes:
    document = Document()
    _build_header(document)
    _build_footer(document)
    _add_page_watermark(document)

    heading = document.add_paragraph()
    h_run = heading.add_run(title)
    h_run.font.bold = True
    h_run.font.size = Pt(16)
    h_run.font.color.rgb = PRIMARY_RGB

    client = f"   ·   Prepared for {client_name}" if client_name else ""
    meta = document.add_paragraph(
        f"Dataset: {dataset_name}   ·   {date.today():%d %B %Y}{client}"
    )
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x6A, 0x70, 0x7A)
    document.add_paragraph()

    if kpis:
        _add_kpi_row(document, kpis)
    _add_summary(document, results)

    for result in results:
        _add_section(document, result)
    _add_signature(document)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
